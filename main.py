# main.py
import subprocess
import os
import datetime
from docx import Document

def run_script(script_name, args=[]):
    script_path = os.path.join("core", script_name)
    try:
        result = subprocess.run(['python3', script_path] + args, capture_output=True, text=True)
        return result.stdout.strip()
    except Exception as e:
        return f"Error running {script_name}: {str(e)}"

layer2_attacks = {
    "1": ("arp_poisoning.py", "ARP Poisoning Attack"),
    "2": ("dns_poisoning.py", "DNS Cache Poisoning Attack"),
    "3": ("cdp_flooding.py", "CDP Flooding Attack"),
    "4": ("mac_flooding.py", "MAC Flooding Attack"),
    "5": ("stp_root_guard.py", "STP Root Guard Attack"),
    "6": ("dhcp_flooding.py", "DHCP Flooding Attack"),
    "7": ("dtp_attack.py", "DTP Attack")
}

layer3_attacks = {
    "1": ("ip_spoofing.py", "IP Spoofing Attack"),
    "2": ("icmp_flood.py", "ICMP Flood Attack"),
    "3": ("udp_flood.py", "UDP Flood Attack"),
    "4": ("tcp_syn_flood.py", "TCP SYN Flood Attack")
}

def save_combined_report(all_reports, filename):
    doc = Document()
    doc.add_heading("Segmentation Testing Report", 0)

    doc.add_paragraph("DISCLAIMER:")
    doc.add_paragraph("This testing is conducted with authorized consent. Attacks performed are part of a controlled environment for security assessment. Unauthorized use is strictly prohibited.\n")

    for report in all_reports:
        doc.add_paragraph("\n--- Attack Report ---", style='Heading 2')
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light Grid'
        for header, value in report.items():
            row_cells = table.add_row().cells
            row_cells[0].text = f"{header}:"
            row_cells[1].text = value

    doc.save(filename)

def get_common_fields(target_ip, source_ip, layer_type, output):
    return {
        "Targeted IP": target_ip,
        "Source IP": source_ip,
        "Method used": "Scapy/Custom",
        "Port Used": "N/A",
        "Vulnerability details": layer_type,
        "Attack Vector": "Low to Medium",
        "Attack Complexity": "Low",
        "Privileges Required": "None",
        "User Interaction": "Local",
        "Scope": "Partial",
        "Confidentiality": "Partial",
        "Integrity": "Partial",
        "Availability": "Medium",
        "Severity-Rating": "Possible lateral movement, traffic capture",
        "Business impact": "Use segmentation, filtering, and monitoring",
        "Remediation": "",
        "Proof of Concept": output
    }

def main():
    print("=== Network Segmentation Testing Toolkit ===")
    print("DISCLAIMER: This tool is intended for authorized security testing only.")
    print("Unauthorized use may be punishable by law.")
    confirm = input("Do you agree to proceed? (yes/no): ").strip().lower()

    if confirm != "yes":
        print("Consent not given. Exiting.")
        return

    print("\nSelect the attack layer:")
    print("1. Layer 2 Attacks")
    print("2. Layer 3 Attacks")
    layer_choice = input("Enter 1 or 2: ").strip()

    if layer_choice == "1":
        selected_attacks = layer2_attacks
        layer_type = "Layer 2 manipulation"
    elif layer_choice == "2":
        selected_attacks = layer3_attacks
        layer_type = "Layer 3 manipulation"
    else:
        print("Invalid choice.")
        return

    print("\nAvailable Attacks:")
    for key, (_, desc) in selected_attacks.items():
        print(f"{key}. {desc}")

    selected = input("Enter attack numbers separated by commas (e.g., 1,3,5): ").strip().split(",")
    target_ip = input("Enter target IP address: ").strip()
    source_ip = input("Enter source/spoofed IP address (if applicable): ").strip()

    all_reports = []

    for choice in selected:
        choice = choice.strip()
        if choice in selected_attacks:
            script_name, desc = selected_attacks[choice]
            confirm_attack = input(f"\nDo you want to execute '{desc}'? (yes/no): ").strip().lower()
            if confirm_attack != "yes":
                print(f"Skipped '{desc}'")
                continue

            print(f"Executing {desc}...")
            output = run_script(script_name, [target_ip, source_ip])
            report_entry = get_common_fields(target_ip, source_ip, layer_type, output)
            all_reports.append(report_entry)
        else:
            print(f"Invalid selection: {choice}")

    if all_reports:
        os.makedirs("output", exist_ok=True)
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        docx_filename = os.path.join("output", f"segmentation_report_{timestamp}.docx")
        save_combined_report(all_reports, docx_filename)
        print(f"\n[+] Report saved to: {docx_filename}")
    else:
        print("No attack executed. Report not generated.")

if __name__ == '__main__':
    main()
